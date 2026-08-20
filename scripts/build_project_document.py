"""Build the comprehensive Word guide for the therapy-switch benchmark.

The document is generated from repository facts and, when present, the local
synthetic quickstart outputs. Run the quickstart first if the evidence appendix
should include its held-out metrics and figures.
"""

from __future__ import annotations

import argparse
import csv
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_URL = "https://github.com/prkharr/advanced-therapy-switch-benchmark"
BUILD_DATE = date(2026, 8, 20)

# standard_business_brief preset, with a named editorial-cover override.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
HEADER_FOOTER_IN = 0.492
CONTENT_WIDTH_IN = 6.5
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "222222"
MUTED = "5E6B78"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "7A5A00"
PALE_GOLD = "FFF8E8"
RED = "9B1C1C"
PALE_RED = "FCE8E6"
GREEN = "216E39"
PALE_GREEN = "E8F5EC"
BORDER = "D9DEE6"


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    name: str = FONT_BODY,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_paragraph_border(paragraph, *, side: str, color: str, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "6")
    border.set(qn("w:color"), color)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, width in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: Sequence[int], *, indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    grid_cols = list(grid)
    while len(grid_cols) < len(widths_dxa):
        grid_col = OxmlElement("w:gridCol")
        grid.append(grid_col)
        grid_cols.append(grid_col)
    for grid_col, width in zip(grid_cols, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        _prevent_row_split(row)
        for cell, width in zip(row.cells, widths_dxa):
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_table_borders(table)


def _style_table_text(table, *, body_size: float = 9.2) -> None:
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=body_size,
                        color=NAVY if r_idx == 0 else INK,
                        bold=True if r_idx == 0 else None,
                    )
        if r_idx == 0:
            _set_repeat_header(row)
            for cell in row.cells:
                _set_cell_shading(cell, LIGHT_GRAY)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_dxa: Sequence[int],
    *,
    body_size: float = 9.2,
    first_col_bold: bool = False,
) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
    set_table_geometry(table, widths_dxa)
    _style_table_text(table, body_size=body_size)
    if first_col_bold:
        for row in table.rows[1:]:
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.color.rgb = _rgb(DARK_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])
    set_run_font(run, size=8.5, color=MUTED)


def _set_numbering_geometry(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, indent])
    level.extend([start, num_fmt, level_text, level_jc, p_pr])
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    indent = p_pr.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        p_pr.append(indent)
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")


def add_list(doc: Document, items: Iterable[str], *, numbered: bool = False) -> None:
    num_id = _set_numbering_geometry(doc, bullet=False) if numbered else doc._number_ids["bullet"]
    for item in items:
        paragraph = doc.add_paragraph()
        _apply_num(paragraph, num_id)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.167
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> object:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, color=NAVY, bold=True)
        remainder = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(remainder, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=INK)
    return paragraph


def add_callout(
    doc: Document,
    title: str,
    text: str,
    *,
    fill: str = CALLOUT,
    title_color: str = DARK_BLUE,
) -> object:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title)
    set_run_font(title_run, size=10.5, color=title_color, bold=True)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_paragraph.paragraph_format.line_spacing = 1.05
    body_run = body_paragraph.add_run(text)
    set_run_font(body_run, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    _set_paragraph_shading(paragraph, LIGHT_GRAY)
    run = paragraph.add_run(text)
    set_run_font(run, name=FONT_MONO, size=8.5, color=NAVY)


def add_figure(doc: Document, path: Path, caption: str, alt_text: str) -> bool:
    if not path.exists():
        return False
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(6.15))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", caption)
    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=MUTED, italic=True)
    return True


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal.font.size = Pt(11)
    normal.font.color.rgb = _rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.color.rgb = _rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT_BODY
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = _rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    if "Section Lead" not in doc.styles:
        lead = doc.styles.add_style("Section Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = doc.styles["Section Lead"]
    lead.font.name = FONT_BODY
    lead._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    lead._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    lead.font.size = Pt(11)
    lead.font.italic = True
    lead.font.color.rgb = _rgb(MUTED)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.10
    lead.paragraph_format.keep_with_next = True
    lead.paragraph_format.keep_together = True


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_IN)
    section.footer_distance = Inches(HEADER_FOOTER_IN)
    section.different_first_page_header_footer = True

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = ""
    header_paragraph.paragraph_format.space_after = Pt(0)
    table = header.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    set_table_geometry(table, [6480, 2880], indent_dxa=0)
    # Header furniture intentionally has no visible borders.
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = table._tbl.tblPr.find(qn("w:tblBorders")).find(qn(f"w:{edge}"))
        node.set(qn("w:val"), "nil")
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left.add_run("ADVANCED THERAPY SWITCH BENCHMARK")
    set_run_font(left_run, size=8.5, color=MUTED, bold=True)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right.add_run("PROJECT GUIDE")
    set_run_font(right_run, size=8.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ""
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    set_table_geometry(footer_table, [7200, 2160], indent_dxa=0)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = footer_table._tbl.tblPr.find(qn("w:tblBorders")).find(qn(f"w:{edge}"))
        node.set(qn("w:val"), "nil")
    left = footer_table.cell(0, 0).paragraphs[0]
    left_run = left.add_run("Commercial analytics - not clinical guidance")
    set_run_font(left_run, size=8.5, color=MUTED)
    right = footer_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = right.add_run("Page ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_number(right)


def add_cover(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(78)
    kicker.paragraph_format.space_after = Pt(18)
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("MACHINE LEARNING BENCHMARKING FRAMEWORK")
    set_run_font(run, size=10.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Advanced Therapy Switch Prediction")
    set_run_font(title_run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle_run = subtitle.add_run("Classical ML vs Deep Learning")
    set_run_font(subtitle_run, size=16, color=DARK_BLUE)

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(32)
    descriptor_run = descriptor.add_run("Complete project architecture, implementation, evaluation, and operating guide")
    set_run_font(descriptor_run, size=11, color=MUTED, italic=True)

    table = doc.add_table(rows=5, cols=2)
    rows = [
        ("Business use", "Commercial patient ranking and HCP opportunity prioritization"),
        ("Research question", "Can sequence-based deep learning materially improve on strong claims baselines?"),
        ("Data modes", "Synthetic development data and canonicalized real claims"),
        ("Repository", "prkharr/advanced-therapy-switch-benchmark"),
        ("Guide date", BUILD_DATE.strftime("%d %B %Y")),
    ]
    for r_idx, (label, value) in enumerate(rows):
        table.cell(r_idx, 0).text = label
        table.cell(r_idx, 1).text = value
    set_table_geometry(table, [2700, 6660])
    for r_idx, row in enumerate(table.rows):
        _set_cell_shading(row.cells[0], BLUE_GRAY)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=NAVY, bold=True)
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    repo = doc.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.paragraph_format.space_after = Pt(20)
    add_hyperlink(repo, "Open the project on GitHub", REPO_URL)

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(10)
    notice.paragraph_format.space_after = Pt(0)
    _set_paragraph_shading(notice, PALE_GOLD)
    notice_run = notice.add_run(
        "Synthetic benchmark results in this guide are engineering evidence only. "
        "They are not a production model claim and must not support treatment decisions."
    )
    set_run_font(notice_run, size=9.5, color=GOLD, bold=True)
    doc.add_page_break()


def add_section_heading(doc: Document, title: str, lead: str, *, new_page: bool = False) -> None:
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.page_break_before = new_page
    lead_paragraph = doc.add_paragraph(style="Section Lead")
    lead_run = lead_paragraph.add_run(lead)
    set_run_font(lead_run, size=11, color=MUTED, italic=True)


def add_document_map(doc: Document) -> None:
    add_section_heading(
        doc,
        "How to use this guide",
        "A practical map for business, analytics, engineering, validation, and deployment reviewers.",
        new_page=False,
    )
    add_body(
        doc,
        "This guide explains what the repository does, why each layer exists, how to run it, and what must change before production claims are introduced. It distinguishes implemented behavior from configurable policy and from work that remains a production gate.",
    )
    rows = [
        ("1-3", "Decision framing", "Business objective, end-to-end flow, strict prediction timeline, and intended use."),
        ("4-6", "Data foundation", "Canonical schemas, synthetic generator, cohort rules, tabular features, and sequences."),
        ("7-9", "Benchmark science", "Model families, controlled ML-vs-DL comparisons, tuning, ranking metrics, uncertainty, and calibration."),
        ("10-12", "Activation and operations", "HCP attribution, opportunity scoring, repository structure, run commands, and outputs."),
        ("13-14", "Evidence and next steps", "Validated synthetic quickstart, limitations, production gates, governance, and roadmap."),
        ("Appendices", "Reference", "Configuration map, important files, output catalog, and glossary."),
    ]
    add_table(doc, ["Sections", "Focus", "What the reader gets"], rows, [1440, 2520, 5400], first_col_bold=True)

    doc.add_heading("Evidence labels", level=2)
    add_table(
        doc,
        ["Label", "Meaning", "How to interpret it"],
        [
            ("Implemented", "Behavior exists in the repository and is exercised by tests or the quickstart.", "Treat as current code behavior."),
            ("Configurable", "A setting or governed rule can change without rewriting model logic.", "Validate the selected value for each data delivery."),
            ("Production gate", "A dependency, validation, or policy decision remains before live scoring.", "Do not infer readiness from synthetic execution alone."),
        ],
        [1800, 3780, 3780],
        first_col_bold=True,
    )


def section_executive_summary(doc: Document) -> None:
    add_section_heading(
        doc,
        "1. Executive summary",
        "The framework is complete enough for controlled engineering benchmarks while remaining deliberately conservative about production claims.",
    )
    add_callout(
        doc,
        "Outcome",
        "The repository provides an end-to-end, modular benchmark from canonical claims through patient scoring and HCP prioritization. It preserves the key scientific distinction between an MLP on aggregate features and true sequence models on event histories.",
        fill=PALE_GREEN,
        title_color=GREEN,
    )
    add_list(
        doc,
        [
            "Two experiments are supported: a patient-stratified random split and a later-patient out-of-time split. The temporal view is the default decision view.",
            "Eleven benchmark rows are always retained: naive baseline, logistic regression, random forest, XGBoost, LightGBM, CatBoost, MLP, LSTM, GRU, BiLSTM, and Transformer.",
            "Unavailable optional libraries are reported as NOT APPLICABLE with reasons. No score is fabricated and no model silently disappears.",
            "Ranking evidence - PR-AUC, capture, lift, deciles, cumulative gains, calibration, confidence intervals, cost, and complexity - matters more than accuracy.",
            "Patient prediction and HCP opportunity ranking are separate, auditable layers. The output is commercial analytics, not a treatment recommendation.",
        ],
    )

    doc.add_heading("Current validation snapshot", level=2)
    add_table(
        doc,
        ["Area", "Current evidence", "Interpretation"],
        [
            ("Automated tests", "40 passed; 1 PyTorch-dependent test skipped locally", "Core, data, leakage, models, evaluation, HCP, and orchestration paths passed."),
            ("Synthetic quickstart", "600 patients; 48 switchers; 8% prevalence; 79 aggregate features", "The full non-neural path executed on a learnable but noisy synthetic population."),
            ("Split coverage", "Stratified and temporal experiments", "Random performance is separated from out-of-time performance."),
            ("Optional models", "Boosting and PyTorch families marked NOT APPLICABLE in the local quickstart", "The framework behaved transparently; this is not evidence that those model families fail."),
            ("CI", "GitHub Actions configured and passing before this document update", "Repository validation is automated on push."),
        ],
        [1980, 3240, 4140],
        first_col_bold=True,
    )


def section_business_objective(doc: Document) -> None:
    add_section_heading(
        doc,
        "2. Business objective and decision framing",
        "The prediction problem is an imbalanced binary classification task whose operating value comes from ranking scarce field capacity.",
    )
    add_body(
        doc,
        "For eligible patient i, the model estimates P(Y_i = 1 | history available at the index date). A positive outcome means an advanced therapy is initiated strictly after the index date and inside the configured prediction window. The resulting probability is stored as advanced_therapy_propensity_score.",
    )

    doc.add_heading("End-to-end business workflow", level=2)
    add_table(
        doc,
        ["Data foundation", "Patient modeling", "Commercial activation"],
        [
            ("Longitudinal claims\nEligible population\nIndex date and history", "Features and event sequences\nPropensity benchmark\nPatient ranking", "Patient-to-HCP attribution\nHCP opportunity score\nField-force prioritization"),
        ],
        [3120, 3120, 3120],
        body_size=9.8,
    )

    doc.add_heading("Primary research question", level=2)
    add_callout(
        doc,
        "Controlled question",
        "Can deep-learning models materially outperform strong traditional machine-learning models for future advanced-therapy switch ranking, after controlling cohort, time boundary, split, evaluation population, and decision metric?",
        fill=BLUE_GRAY,
    )

    doc.add_heading("Intended and prohibited use", level=2)
    add_table(
        doc,
        ["Intended use", "Prohibited use"],
        [
            ("Rank eligible patients by predictive association with a future advanced-therapy claim.", "Treatment selection, diagnosis, medical advice, or clinical decision support."),
            ("Aggregate probabilities into transparent HCP opportunity metrics for commercial review.", "Causal claims about a patient, provider, therapy, diagnosis, or feature."),
            ("Compare model families and decide whether added complexity is justified.", "Automated adverse decisions or use outside approved populations, fields, and geographies."),
        ],
        [4680, 4680],
    )


def section_temporal_design(doc: Document) -> None:
    add_section_heading(
        doc,
        "3. Prediction timeline and leakage control",
        "Time boundaries are enforced as pipeline invariants, not left to analyst convention.",
    )
    add_table(
        doc,
        ["Observation window", "Index date", "Prediction window"],
        [
            ("[index - 365 days, index]\nFeatures and sequence events only", "Configurable conventional-claim anchor\nAll information frozen here", "(index, index + 90 days]\nOutcome only"),
        ],
        [3900, 1560, 3900],
        body_size=9.8,
    )
    add_body(
        doc,
        "The 365-day lookback and 90-day prediction horizon are defaults, not hard-coded assumptions. The scoring contract permits index-date events only when they would be operationally available by score time. If claim latency makes that unrealistic, a data-lag cutoff must be added and tested.",
    )

    doc.add_heading("Automated controls", level=2)
    add_list(
        doc,
        [
            "Cohort logic excludes advanced therapy at or before index and labels only advanced fills strictly after index.",
            "Feature and sequence builders join each event to a patient-specific index and discard post-index or out-of-lookback events.",
            "Structural audits reject forbidden outcome-like columns, duplicate patient snapshots, infinite values, label mismatches, and unknown patients.",
            "Split functions assert that no patient appears in more than one partition; temporal splitting keeps equal index dates together.",
            "Preprocessing, sequence vocabularies, hyperparameters, early stopping, threshold choice, and calibration are learned without the test population.",
            "HCP attribution uses only pre-index interactions. Future prescribers or visits are invalid evidence.",
        ],
    )
    add_callout(
        doc,
        "Test coverage",
        "tests/test_leakage.py injects distinctive future claims and verifies aggregate-feature invariance, rejects explicit post-index events, checks forbidden columns, and validates patient-disjoint splits. Sequence tests separately confirm that all retained event dates are on or before index.",
    )


def section_data_architecture(doc: Document) -> None:
    add_section_heading(
        doc,
        "4. Data architecture and real-claims adaptation",
        "Source-specific fields are normalized once; all downstream cohort, feature, model, and HCP code operates on a vendor-neutral contract.",
    )
    add_table(
        doc,
        ["Canonical table", "Required content", "Role"],
        [
            ("patients", "patient_id, gender, geography, observation_start/end; age or birth_year", "Eligibility, demographics, and observable coverage."),
            ("medical_claims", "claim_id, patient_id, claim_date, diagnosis/procedure codes, provider, place of service", "Disease burden, utilization, procedures, specialists, and event sequences."),
            ("pharmacy_claims", "claim_id, patient_id, fill_date, drug_id, therapy_class, quantity, days_supply, prescriber", "Indexing, outcome, therapy history, adherence proxies, attribution, and sequences."),
            ("providers", "provider_id, specialty, geography, organization", "Specialist features and HCP-level activation."),
        ],
        [1800, 4380, 3180],
        first_col_bold=True,
    )

    doc.add_heading("Adapter pattern", level=2)
    add_body(
        doc,
        "Set data.source to files, select CSV or Parquet, point data.input_dir at the governed snapshot, and map delivery-specific names under data.tables.<table>.columns. The loader validates table presence, required columns, date parseability, and non-null patient identifiers.",
    )
    add_code_block(
        doc,
        "data:\n"
        "  source: files\n"
        "  input_dir: D:/governed/claims_snapshot\n"
        "  file_format: parquet\n"
        "  tables:\n"
        "    pharmacy_claims:\n"
        "      file: rx_claims\n"
        "      columns:\n"
        "        service_date: fill_date\n"
        "        product_code: drug_id\n"
        "        mapped_class: therapy_class",
    )
    add_callout(
        doc,
        "Governed mappings",
        "Conventional and advanced therapy definitions belong in environment configuration, not modeling logic. Production owners must version them, review effective dates, and ensure that the two sets do not overlap.",
        fill=PALE_GOLD,
        title_color=GOLD,
    )

    doc.add_heading("Production data checks beyond the generic contract", level=2)
    add_list(
        doc,
        [
            "Adjudication status, reversals, duplicate lines, code systems, enrollment completeness, and claim run-out.",
            "Cash fills, assistance channels, specialty-pharmacy feeds, and other sources that alter label observability.",
            "Claim latency and the operational as-of timestamp used for scoring.",
            "Permitted demographics and geography, tokenization rules, access control, and retention policy.",
        ],
    )


def section_synthetic_data(doc: Document) -> None:
    add_section_heading(
        doc,
        "5. Synthetic claims generator",
        "The synthetic mode enables reproducible end-to-end development without copying, inferring, or reverse-engineering proprietary claims data.",
    )
    add_body(
        doc,
        "src/therapy_switch/data/generate_synthetic_claims.py creates patients, providers, medical and pharmacy claims, diagnoses, procedures, specialties, conventional exposure, advanced initiation, and longitudinal utilization. It returns the same four canonical tables used by the real-data path.",
    )

    doc.add_heading("Signal design", level=2)
    add_table(
        doc,
        ["Pattern", "Synthetic manifestation", "Why it matters"],
        [
            ("Disease escalation", "Higher recent claim intensity, severe diagnosis proxies, and procedures", "Creates learnable deterioration without exposing the label."),
            ("Specialist involvement", "Specialist affinity and more recent relevant specialist activity", "Represents care-pathway escalation."),
            ("Treatment history", "Refill cadence, conventional duration, therapy changes, and failure-like patterns", "Supports both static history and temporal modeling."),
            ("Noise and imbalance", "Configurable prevalence, Gaussian/Gumbel noise, latent variability, and controls switching beyond horizon", "Prevents trivial separation and keeps ranking realistic."),
            ("Leakage challenge", "Raw post-index utilization remains in generated tables", "Exercises feature and sequence filters rather than relying on clean input."),
        ],
        [1800, 3960, 3600],
        first_col_bold=True,
    )
    add_body(
        doc,
        "The default engineering population contains 5,000 patients and 350 providers with an 8% target prevalence. The quickstart reduces this to 600 patients and 70 providers for fast smoke testing. Signal strength, noise scale, date span, prevalence, and random seed are configurable.",
    )
    add_callout(
        doc,
        "Boundary of evidence",
        "Synthetic performance proves that the pipeline can learn, rank, evaluate, and export. It does not estimate production accuracy, market lift, patient behavior, or Komodo Healthcare Map performance.",
        fill=PALE_RED,
        title_color=RED,
    )


def section_cohort_features(doc: Document) -> None:
    add_section_heading(
        doc,
        "6. Cohort, aggregate features, and event sequences",
        "One cohort supports two model representations: a patient-level feature table and an auditable pre-index event tensor.",
    )
    doc.add_heading("Cohort construction", level=2)
    add_list(
        doc,
        [
            "Require sufficient observable lookback and future follow-up for each patient.",
            "Anchor an index date using a configured conventional-therapy claim strategy.",
            "Require conventional exposure during the lookback and, by default, on the index date.",
            "Exclude advanced therapy at or before index.",
            "Set label 1 when the earliest advanced fill falls in (index, prediction end]; otherwise set label 0.",
        ],
        numbered=True,
    )

    doc.add_heading("Aggregate feature families", level=2)
    add_table(
        doc,
        ["Feature family", "Representative variables"],
        [
            ("Demographics", "Age, gender, and permitted geography."),
            ("Disease burden", "Diagnosis count, unique diagnoses, comorbidity count, severity proxies, and configured disease flags."),
            ("Treatment history", "Current therapy, prior therapies, observed duration, changes, refill frequency, PDC, and discontinuation proxy."),
            ("Utilization", "Outpatient, inpatient, ER, specialist, procedure, pharmacy, and visit counts."),
            ("Recency", "Days since last diagnosis, RX, specialist visit, and treatment change."),
            ("Trend and velocity", "Recent-versus-prior ratios, diagnosis change, RX-frequency change, specialist acceleration, and utilization trend."),
        ],
        [2160, 7200],
        first_col_bold=True,
    )
    add_body(
        doc,
        "Counts are calculated over configurable 30, 60, 90, 180, and 365-day windows. Ratios use safe smoothing, missing recency receives a configurable sentinel, and adherence logic unions overlapping supply intervals.",
    )

    doc.add_heading("Sequence representation", level=2)
    add_body(
        doc,
        "Medical rows become diagnosis and procedure events; pharmacy fills become medication events. After chronological sorting and lookback filtering, the most recent max_length events are retained and right-padded. Training vocabularies are reused for validation and test so unseen categories map to <UNK>.",
    )
    add_table(
        doc,
        ["Event field", "Representation", "Audit value"],
        [
            ("Event type", "Diagnosis, procedure, or pharmacy token", "Distinguishes care actions."),
            ("Code", "Diagnosis, procedure, or product token", "Retains granular event identity."),
            ("Therapy class", "Conventional, advanced, unknown, or padding token", "Supports treatment trajectory."),
            ("Provider specialty", "Pre-index specialty token", "Adds care-context signal."),
            ("Time", "Time since previous event and days before index", "Captures spacing, recency, and ordering."),
            ("Mask and dates", "Attention mask plus retained event/index dates", "Enables padding control and a second leakage audit."),
        ],
        [1980, 3420, 3960],
        first_col_bold=True,
    )


def section_models(doc: Document) -> None:
    add_section_heading(
        doc,
        "7. Model catalog and controlled comparisons",
        "The benchmark asks two different questions and prevents a tabular MLP result from standing in for all deep learning.",
    )
    doc.add_heading("Comparison A: tabular", level=2)
    add_body(
        doc,
        "Logistic regression, random forest, XGBoost, LightGBM, CatBoost, and MLP receive the same leakage-safe aggregate feature frame and the same split. This isolates algorithm choice on structured data.",
    )
    doc.add_heading("Comparison B: longitudinal", level=2)
    add_body(
        doc,
        "The validation-selected best aggregate classical model is compared with LSTM, GRU, BiLSTM, and a compact temporal Transformer built from pre-index event sequences. Bidirectionality is permitted only because both directions traverse a completed historical window with no post-index event.",
    )
    add_callout(
        doc,
        "Critical interpretation rule",
        "If the MLP loses to gradient boosting, the correct conclusion is about tabular neural modeling on aggregate features. Longitudinal deep learning must be judged separately when valid event sequences and dependencies are available.",
        fill=PALE_GOLD,
        title_color=GOLD,
    )

    doc.add_heading("Implemented model registry", level=2)
    rows = [
        ("Naive Baseline", "Baseline", "Constant training prevalence", "Random-targeting reference"),
        ("Logistic Regression", "Classical", "Scaled numeric + one-hot categorical", "Balanced class weights"),
        ("Random Forest", "Classical", "Aggregate tabular", "Balanced subsample weights"),
        ("XGBoost", "Classical", "Aggregate tabular", "scale_pos_weight + early stopping"),
        ("LightGBM", "Classical", "Aggregate tabular", "scale_pos_weight + early stopping"),
        ("CatBoost", "Classical", "Aggregate tabular", "scale_pos_weight + early stopping"),
        ("MLP", "Tabular DL", "Same aggregate tabular frame", "Class weights; BCE vs focal; dropout"),
        ("LSTM", "Longitudinal DL", "Pre-index event tensor", "Weighted neural loss + early stopping"),
        ("GRU", "Longitudinal DL", "Pre-index event tensor", "Weighted neural loss + early stopping"),
        ("BiLSTM", "Longitudinal DL", "Historical events in both directions", "Second sequence leakage audit"),
        ("Transformer", "Longitudinal DL", "Event projection, position/time encoding, attention", "Masked pooling + compact head"),
    ]
    add_table(doc, ["Model", "Category", "Representation", "Imbalance/control"], rows, [1980, 1620, 3240, 2520], body_size=8.5, first_col_bold=True)


def section_training(doc: Document) -> None:
    add_section_heading(
        doc,
        "8. Training, imbalance, tuning, and calibration",
        "Every optimization choice is fit on training or validation data; the final test population remains untouched until evaluation.",
    )
    doc.add_heading("Data splitting", level=2)
    add_table(
        doc,
        ["Experiment", "Assignment", "Use"],
        [
            ("Patient-stratified", "Random patient-disjoint train/validation/test partitions with approximate class balance", "Development sensitivity and comparability with conventional benchmarks."),
            ("Temporal / out-of-time", "Earlier index dates to train, later dates to validation, latest dates to test", "Primary realism check for changing care and coding patterns."),
        ],
        [2160, 3600, 3600],
        first_col_bold=True,
    )
    add_body(doc, "The default fractions are 65% train, 15% validation, and 20% test. Patients sharing an index date stay together in temporal splitting.")

    doc.add_heading("Imbalance handling", level=2)
    add_list(
        doc,
        [
            "Report positive count, negative count, prevalence, and positive-to-negative ratio before training.",
            "Use balanced weights for logistic regression and random forest, scale_pos_weight for supported boosters, and positive-class weights for neural losses.",
            "Compare binary cross-entropy and focal loss for the MLP using validation PR-AUC.",
            "Tune the operating threshold on validation data with max-F1 or a fixed threshold. Do not use accuracy as the optimization target.",
            "SMOTE is not used by the current framework. Any future oversampling must stay inside the training partition and tabular pipeline only.",
        ],
    )

    doc.add_heading("Hyperparameter optimization", level=2)
    add_body(
        doc,
        "When tuning is enabled, Optuna is preferred; otherwise RandomizedSearchCV is the fallback. The objective is average precision / PR-AUC. The framework records the best parameters, validation score, backend, duration, and trial count. Current orchestration applies automated tuning to classical estimators; neural architecture search remains configuration-driven.",
    )

    doc.add_heading("Calibration", level=2)
    add_body(
        doc,
        "Raw, Platt/sigmoid, and isotonic candidates are fitted on validation predictions and compared by validation Brier score. The selected calibrator is then applied once to test predictions. Isotonic calibration should be rejected when validation positives are too sparse to support a stable monotonic mapping.",
    )


def section_evaluation(doc: Document) -> None:
    add_section_heading(
        doc,
        "9. Evaluation, uncertainty, and model decision policy",
        "The project measures both statistical discrimination and the practical value of targeting a limited population fraction.",
    )
    doc.add_heading("Metric hierarchy", level=2)
    add_table(
        doc,
        ["Layer", "Metrics", "Decision value"],
        [
            ("Ranking - primary", "PR-AUC, Recall/Precision/Lift at top 1/5/10/20/30%, deciles, cumulative gains", "How effectively the model concentrates future switchers within field capacity."),
            ("Discrimination", "ROC-AUC", "Overall pairwise ranking quality; secondary under strong imbalance."),
            ("Operating point", "Precision, recall, specificity, F1, balanced accuracy", "Behavior at the validation-selected threshold."),
            ("Probability quality", "Log loss, Brier score, calibration curve", "Whether scores behave like probabilities."),
            ("Uncertainty and stability", "95% bootstrap intervals, paired bootstrap, random-vs-OOT deltas", "Whether small differences are credible and stable."),
            ("Operational burden", "Training time, inference time, explainability, implementation complexity", "Whether added performance justifies added cost."),
        ],
        [1800, 3960, 3600],
        body_size=8.8,
        first_col_bold=True,
    )

    doc.add_heading("Capacity metrics", level=2)
    add_code_block(
        doc,
        "Recall@TopX = switchers in highest-scoring X% / all switchers\n"
        "Precision@TopX = switchers in highest-scoring X% / patients targeted\n"
        "Lift@TopX = switch rate in highest-scoring X% / overall switch rate",
    )
    add_body(
        doc,
        "Decile 1 is the highest-propensity tenth of the test population. The decile table records patient count, switchers, rate, precision, recall in decile, cumulative recall, lift, cumulative lift, and mean predicted probability for every completed model.",
    )

    doc.add_heading("Confidence and recommendation", level=2)
    add_list(
        doc,
        [
            "Bootstrap the held-out population for ROC-AUC, PR-AUC, Recall@10/20%, and Lift@10/20%.",
            "Use paired resamples when the best classical and best sequence model score the same patients.",
            "Require uncertainty, out-of-time stability, field-capacity metrics, calibration, and cost to support a material improvement.",
            "Use the configured 0.01 PR-AUC materiality threshold as a policy input, not as a substitute for confidence intervals and operating context.",
            "Prefer the simplest adequate model when added complexity is not credibly justified; do not generalize one dataset's result to all deep learning.",
        ],
    )

    doc.add_heading("Explainability", level=2)
    add_body(
        doc,
        "Tree models use SHAP when available; deterministic native or permutation importance is the fallback. Outputs include global importance and top positive/negative drivers. Generic tabular attribution is explicitly rejected for sequence tensors; a sequence-specific method such as integrated gradients is a production gate. All explanations describe predictive association, not causality.",
    )


def section_hcp(doc: Document) -> None:
    add_section_heading(
        doc,
        "10. HCP attribution and opportunity prioritization",
        "Patient prediction is preserved as a separate analytical layer from HCP attribution and commercial ranking.",
    )
    doc.add_heading("Attribution rules", level=2)
    add_table(
        doc,
        ["Rule", "Definition", "Governance note"],
        [
            ("Most recent relevant prescriber", "Latest pre-index relevant prescription interaction", "Current default; replace when a governed business rule already exists."),
            ("Most frequent relevant specialist", "Provider with the greatest relevant visit frequency", "Tie-breaking and specialty eligibility must be documented."),
            ("Plurality of relevant claims", "Provider with the largest count of configured relevant claims", "Requires an agreed claim universe."),
            ("Existing attribution", "Use an externally supplied patient-to-HCP mapping", "Preferred when the organization already governs attribution."),
        ],
        [2520, 3420, 3420],
        first_col_bold=True,
    )

    doc.add_heading("Opportunity metrics", level=2)
    add_body(
        doc,
        "For each HCP, the framework reports eligible_patient_count, high_propensity_patient_count, patients_top_5pct, patients_top_10pct, patients_top_20pct, mean_patient_propensity, max_patient_propensity, and expected_switchers.",
    )
    add_code_block(doc, "expected_switchers = SUM(patient advanced_therapy_propensity_score)")

    doc.add_heading("Transparent score", level=2)
    add_body(
        doc,
        "The default opportunity score is a weighted sum of percentile-normalized HCP components: 60% expected switchers, 25% high-propensity patients, and 15% eligible-patient volume. The weights are configurable business rules, not learned treatment effects.",
    )
    add_body(
        doc,
        "Exact top-5/10/20% bands use a fixed patient count. The separate high-propensity threshold can include more than its nominal percentile when calibrated scores tie at the boundary; this behavior should be reported to field-planning users.",
    )
    add_callout(
        doc,
        "Separation of responsibilities",
        "The patient model estimates future switch propensity. The HCP layer maps and aggregates those estimates. A change in attribution or field capacity should not require retraining the patient model unless the prediction target itself changes.",
        fill=BLUE_GRAY,
    )


def section_implementation(doc: Document) -> None:
    add_section_heading(
        doc,
        "11. Repository structure and execution",
        "The project is a reusable Python package with configuration, CLI entry points, tests, documentation, outputs, and model artifacts.",
    )
    add_table(
        doc,
        ["Path", "Responsibility"],
        [
            ("configs/", "Run settings, temporal windows, model parameters, evaluation policy, and non-proprietary therapy mappings."),
            ("src/therapy_switch/data/", "Synthetic generation, cohort construction, splitting, and event sequences."),
            ("src/therapy_switch/features/", "Aggregate feature engineering and leakage audits."),
            ("src/therapy_switch/models/", "Contracts, preprocessing, model registry, classical/neural runners, architectures, and tuning."),
            ("src/therapy_switch/evaluation/", "Metrics, deciles/gains, bootstrap, calibration, explainability, and plots."),
            ("src/therapy_switch/hcp/", "Patient-to-HCP attribution and opportunity ranking."),
            ("src/therapy_switch/pipeline.py", "End-to-end orchestration across experiments and full-cohort scoring."),
            ("tests/", "Unit/integration coverage including explicit temporal leakage tests."),
            ("docs/", "Data contract, leakage strategy, benchmark design, model-card template, and production checklist."),
        ],
        [3240, 6120],
        body_size=9.0,
        first_col_bold=True,
    )

    doc.add_heading("Quickstart", level=2)
    add_code_block(
        doc,
        "py -3.10 -m venv .venv\n"
        ".venv\\Scripts\\Activate.ps1\n"
        "python -m pip install -e \".[dev]\"\n"
        "therapy-switch run --config configs/quickstart.yaml\n"
        "pytest",
    )
    doc.add_heading("Full optional model suite", level=2)
    add_code_block(
        doc,
        "python -m pip install -e \".[all,dev]\"\n"
        "therapy-switch run --config configs/default.yaml",
    )
    doc.add_heading("Useful CLI operations", level=2)
    add_code_block(
        doc,
        "therapy-switch generate --config configs/quickstart.yaml --output-dir data/synthetic\n"
        "therapy-switch validate-data --config configs/real_claims.yaml\n"
        "therapy-switch run --config configs/default.yaml --experiment both\n"
        "therapy-switch run --config configs/default.yaml --experiment temporal --no-plots",
    )
    add_body(
        doc,
        "Core execution depends on scikit-learn. Optional packages activate XGBoost, LightGBM, CatBoost, PyTorch neural models, Optuna tuning, SHAP explanations, or Parquet I/O. Missing model dependencies remain visible in benchmark outputs with a reason.",
    )


def section_outputs(doc: Document) -> None:
    add_section_heading(
        doc,
        "12. Outputs, artifacts, and reproducibility",
        "Human-readable benchmark evidence is separated from fitted model artifacts and sensitive row-level production data.",
    )
    add_table(
        doc,
        ["Output", "Purpose"],
        [
            ("model_benchmark.csv", "Full 11-model table with ranking, classification, calibration, timing, status, and reason columns."),
            ("executive_benchmark.csv", "Presentation-ready PR-AUC, ROC-AUC, capture, lift, complexity, and recommendation view."),
            ("tabular_comparison.csv", "Controlled aggregate-feature model comparison."),
            ("longitudinal_comparison.csv", "Best aggregate classical model versus valid sequence models."),
            ("decile_analysis.csv", "Per-model decile and cumulative ranking performance."),
            ("cumulative_gains.csv", "Population targeted versus future switchers captured."),
            ("bootstrap_confidence_intervals.csv", "Held-out 95% confidence intervals for required metrics."),
            ("paired_model_comparison.csv", "Paired best-classical versus best-DL differences when both are valid."),
            ("random_vs_temporal_benchmark.csv / model_stability.csv", "Experiment labels and out-of-time stability deltas."),
            ("patient_propensity_scores.csv", "Patient-level advanced_therapy_propensity_score."),
            ("hcp_targeting_output.csv", "Attributed HCP opportunity metrics and transparent score."),
            ("model_recommendation.md", "Decision narrative with evidence scope and guardrails."),
        ],
        [3600, 5760],
        body_size=8.7,
        first_col_bold=True,
    )
    add_body(
        doc,
        "Experiment-specific results live under outputs/experiments/<stratified|temporal>/. The configured primary experiment is copied to the top-level outputs directory. Figures include ROC/PR, calibration, lift/gains, deciles, top-fraction lift, and benchmark comparisons.",
    )
    doc.add_heading("Artifact trail", level=2)
    add_list(
        doc,
        [
            "Resolved configuration and fixed random seed.",
            "Run manifest with code/runtime context, cohort, split, features, timings, and status evidence.",
            "Model summaries, serialized estimators, neural training histories, and sequence vocabularies.",
            "Selected thresholds, calibration method, failure reasons, and explainability status.",
        ],
    )
    add_callout(
        doc,
        "Repository hygiene",
        "Generated claims, model binaries, row-level scores, and operational logs are ignored by Git. In production they belong in approved governed storage, not source control.",
        fill=PALE_RED,
        title_color=RED,
    )


def _read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _fmt_metric(value: str) -> str:
    if value is None or value == "":
        return "-"
    try:
        return f"{float(value):.3f}"
    except ValueError:
        return value


def section_quickstart_evidence(doc: Document, repo_root: Path) -> None:
    add_section_heading(
        doc,
        "13. Validated synthetic quickstart evidence",
        "These results document engineering execution on the local synthetic configuration; they are not a forecast of production performance.",
    )
    summary_rows = _read_csv(repo_root / "outputs" / "quickstart" / "cohort_summary.csv")
    benchmark_rows = _read_csv(repo_root / "outputs" / "quickstart" / "model_benchmark.csv")
    if not summary_rows or not benchmark_rows:
        add_callout(
            doc,
            "Quickstart outputs not found",
            "Run therapy-switch run --config configs/quickstart.yaml, then rebuild this guide to populate the evidence table and figures.",
            fill=PALE_GOLD,
            title_color=GOLD,
        )
        return

    summary = summary_rows[0]
    add_table(
        doc,
        ["Patients", "Switchers", "Prevalence", "Index range", "Aggregate features"],
        [
            (
                str(int(float(summary["positive_patients"])) + int(float(summary["negative_patients"]))),
                summary["positive_patients"],
                f"{float(summary['positive_prevalence']) * 100:.1f}%",
                f"{summary['index_date_min']} to {summary['index_date_max']}",
                summary["feature_count"],
            )
        ],
        [1200, 1200, 1440, 3720, 1800],
        body_size=9.0,
    )

    rows = []
    for row in benchmark_rows:
        rows.append(
            (
                row["Model"],
                row["Status"],
                _fmt_metric(row.get("PR-AUC", "")),
                _fmt_metric(row.get("ROC-AUC", "")),
                _fmt_metric(row.get("Recall@10%", "")),
                _fmt_metric(row.get("Lift@10%", "")),
            )
        )
    doc.add_heading("Primary temporal test benchmark", level=2)
    add_table(
        doc,
        ["Model", "Status", "PR-AUC", "ROC-AUC", "Recall@10%", "Lift@10%"],
        rows,
        [2340, 1980, 1260, 1260, 1260, 1260],
        body_size=8.4,
        first_col_bold=True,
    )
    add_body(
        doc,
        "Random Forest was selected on validation PR-AUC (0.111 versus 0.095 for logistic regression), not on the test set. Its temporal-test PR-AUC was 0.285, while logistic regression produced higher ROC-AUC and top-decile capture. Wide bootstrap intervals and only 13 positives in the temporal test subset make these differences uncertain; no production recommendation follows.",
    )
    add_callout(
        doc,
        "What NOT APPLICABLE means here",
        "XGBoost, LightGBM, and CatBoost were not installed in the local environment. PyTorch was also absent, so MLP, LSTM, GRU, BiLSTM, and Transformer did not run. Their blank metrics are explicit dependency evidence, not negative model evidence.",
        fill=PALE_GOLD,
        title_color=GOLD,
    )

    figures = repo_root / "outputs" / "quickstart" / "experiments" / "temporal" / "figures"
    doc.add_heading("Synthetic held-out visuals", level=2)
    add_figure(
        doc,
        figures / "model_benchmark.png",
        "Figure 1. Completed-model benchmark on the synthetic temporal test subset.",
        "Four-panel horizontal bar chart comparing PR-AUC, ROC-AUC, Recall at top 10 percent, and Lift at top 10 percent for the naive baseline, logistic regression, and random forest.",
    )
    add_figure(
        doc,
        figures / "lift_and_gains.png",
        "Figure 2. Cumulative gains and lift on the synthetic temporal test subset.",
        "Two-panel line chart showing percentage of actual switchers captured and cumulative lift as the targeted population fraction increases for the completed models.",
    )


def section_limitations(doc: Document) -> None:
    add_section_heading(
        doc,
        "14. Limitations, production gates, and roadmap",
        "The repository is a strong benchmark foundation, but a live claims program still requires governed data semantics, model validation, and deployment controls.",
    )
    doc.add_heading("Known implementation and documentation gaps", level=2)
    add_table(
        doc,
        ["Gap", "Risk", "Recommended action"],
        [
            ("Some declared settings are not consumed independently", "A reader may assume minimum_history/followup, disease prefixes, sequence min-events/vocab frequency, provider-specialty toggle, or temporal gap already alters behavior", "Wire each setting into execution or remove/document it before production configuration is approved."),
            ("Age precedence differs from the prose contract", "The implementation derives age from birth_year first when both exist", "Select one policy, add a schema rule, and test it."),
            ("Sequence categorical IDs enter a numeric event projection", "Ordinal-looking token IDs may be a weak production representation for high-cardinality claims", "Benchmark dedicated categorical embedding tables and richer code hierarchies."),
            ("Sequence-specific neural attribution is unavailable", "Neural explanations would be incomplete", "Add and validate integrated gradients or another approved sequence method."),
            ("Target-recall threshold strategy is not fully orchestrated", "Configuration may not produce the intended threshold path", "Wire the existing threshold utility and add an end-to-end test before relying on it."),
            ("Automated tuning is classical-only", "Neural comparison may understate a well-tuned sequence model", "Add bounded neural tuning under a fixed compute budget."),
            ("Full-cohort quickstart scoring includes development patients", "It is useful for workflow testing but not prospective evidence", "Score a fresh as-of cohort for deployment validation."),
            ("Patient score exports retain the development label", "A production extract could expose unavailable outcomes or encourage retrospective interpretation", "Remove labels from production scoring exports and retain them only in governed evaluation data."),
            ("Future-claim mutation is not duplicated for sequence invariance", "Existing date assertions are strong but not the same regression pattern", "Add a sequence mutation-invariance test matching the tabular leakage test."),
        ],
        [2880, 3060, 3420],
        body_size=8.2,
        first_col_bold=True,
    )

    doc.add_heading("Production readiness sequence", level=2)
    add_list(
        doc,
        [
            "Approve the business target, eligible population, index-date semantics, claim-lag rule, therapy mappings, and attribution policy.",
            "Build and validate the real-data adapter, including run-out, reversals, channel gaps, coverage, schema, and point-in-time observability.",
            "Execute both split experiments with all permitted dependencies and fixed compute budgets; retain failure rows and training histories.",
            "Review discrimination, capacity metrics, calibration, bootstrap uncertainty, subgroup stability, explainability, and random-vs-OOT drift.",
            "Benchmark inference throughput and storage, finalize monitoring thresholds, rollback, retention, and human approval procedures.",
            "Complete the model card and document the simplest model that meets the business threshold; require explicit go/no-go approval.",
        ],
        numbered=True,
    )

    doc.add_heading("Monitoring after launch", level=2)
    add_list(
        doc,
        [
            "Schema, volume, missingness, therapy-map coverage, index-date distribution, prevalence, and score distribution.",
            "PR-AUC, capture, lift, calibration, and HCP concentration after labels mature.",
            "Subgroup stability for approved demographic and geographic fields, with privacy and fairness review.",
            "Data freshness, claim latency, inference time, error rates, model/version lineage, and attribution coverage.",
        ],
    )
    add_callout(
        doc,
        "Final governance rule",
        "Do not allow model complexity to substitute for evidence. A sequence model should be deployed only when its out-of-time improvement is material, statistically credible, operationally affordable, and explainable enough for the approved commercial use.",
        fill=PALE_GREEN,
        title_color=GREEN,
    )


def appendix_reference(doc: Document) -> None:
    add_section_heading(
        doc,
        "Appendix A. Configuration and code reference",
        "A compact crosswalk from project concern to the primary implementation and documentation paths.",
        new_page=True,
    )
    add_table(
        doc,
        ["Concern", "Primary paths"],
        [
            ("Run configuration", "configs/default.yaml; configs/quickstart.yaml; src/therapy_switch/config.py"),
            ("Schemas and ingestion", "src/therapy_switch/schemas.py; src/therapy_switch/io.py; docs/data_contract.md"),
            ("Synthetic claims", "src/therapy_switch/data/generate_synthetic_claims.py"),
            ("Cohort and labels", "src/therapy_switch/data/cohort.py"),
            ("Aggregate features", "src/therapy_switch/features/feature_engineering.py"),
            ("Event sequences", "src/therapy_switch/data/event_sequences.py"),
            ("Leakage controls", "src/therapy_switch/features/leakage.py; docs/leakage_controls.md; tests/test_leakage.py"),
            ("Splitting", "src/therapy_switch/data/splitting.py"),
            ("Model registry and contracts", "src/therapy_switch/models/registry.py; src/therapy_switch/models/contracts.py"),
            ("Classical models and tuning", "src/therapy_switch/models/classical.py; src/therapy_switch/models/tuning.py"),
            ("Neural models", "src/therapy_switch/models/neural.py; src/therapy_switch/models/architectures.py"),
            ("Evaluation and plots", "src/therapy_switch/evaluation/"),
            ("HCP layer", "src/therapy_switch/hcp/hcp_prioritization.py"),
            ("Orchestration and CLI", "src/therapy_switch/pipeline.py; src/therapy_switch/cli.py"),
            ("Governance", "docs/benchmark_design.md; docs/model_card.md; docs/productionization.md"),
        ],
        [2520, 6840],
        body_size=8.5,
        first_col_bold=True,
    )

    doc.add_heading("Important defaults", level=2)
    add_table(
        doc,
        ["Setting", "Default", "Why configurable"],
        [
            ("Observation / prediction", "365 / 90 days", "Business horizon, data latency, and disease pathway differ."),
            ("Synthetic prevalence", "8%", "Used for engineering imbalance, not a production prior."),
            ("Primary experiment", "Temporal", "Out-of-time evidence is closer to prospective use."),
            ("Train / validation / test", "65% / 15% / 20%", "Population size and positive count may require adjustment."),
            ("Sequence length", "128 events (64 quickstart)", "Compute, truncation, and history density trade off."),
            ("Bootstrap", "1,000 iterations (100 quickstart)", "Final uncertainty needs greater precision than a smoke test."),
            ("HCP score weights", "0.60 / 0.25 / 0.15", "Expected switchers, high-propensity volume, and eligible volume are business choices."),
            ("DL materiality", "0.01 PR-AUC", "Threshold must be interpreted with confidence, stability, cost, and capacity metrics."),
        ],
        [2520, 2340, 4500],
        body_size=8.7,
        first_col_bold=True,
    )


def appendix_glossary(doc: Document) -> None:
    add_section_heading(
        doc,
        "Appendix B. Glossary",
        "Terms used consistently across the code, outputs, and this guide.",
    )
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ("Advanced therapy propensity", "Predicted probability associated with an eligible patient's advanced-therapy initiation during the future window."),
            ("Index date", "Patient-specific cutoff separating allowable history from the future label window."),
            ("Observation window", "Historical period used for eligibility, aggregate features, and event sequences."),
            ("Prediction window", "Strictly post-index period in which advanced initiation defines the outcome."),
            ("PR-AUC", "Average precision; emphasizes positive ranking performance under class imbalance."),
            ("Recall@TopX", "Share of all future switchers contained in the highest-scoring X% of patients."),
            ("Lift@TopX", "Top-X switch rate divided by the overall switch rate."),
            ("Out-of-time", "Evaluation on the most recent index dates, after training and validation on earlier dates."),
            ("Calibration", "Agreement between predicted probabilities and observed event rates."),
            ("Expected switchers", "Sum of patient probabilities attributed to an HCP."),
            ("NOT APPLICABLE", "A model did not validly run because a dependency, input representation, or applicability condition was absent; no result is fabricated."),
        ],
        [2520, 6840],
        body_size=9.0,
        first_col_bold=True,
    )
    add_body(doc, "Repository: ")
    paragraph = doc.paragraphs[-1]
    add_hyperlink(paragraph, "View the repository on GitHub", REPO_URL)
    add_callout(
        doc,
        "Document scope",
        "This guide explains the repository state as validated on 20 August 2026. The YAML configuration, run manifest, model outputs, and production approval record remain the authoritative evidence for any future real-data run.",
        fill=BLUE_GRAY,
    )


def build_document(repo_root: Path, output_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc._number_ids = {
        "bullet": _set_numbering_geometry(doc, bullet=True),
        "number": _set_numbering_geometry(doc, bullet=False),
    }

    props = doc.core_properties
    props.title = "Advanced Therapy Switch Prediction - Complete Project Guide"
    props.subject = "Classical machine learning and longitudinal deep learning benchmark for claims analytics"
    props.keywords = "claims analytics, therapy switch, machine learning, deep learning, HCP prioritization"
    props.creator = ""
    props.last_modified_by = ""

    add_cover(doc)
    add_document_map(doc)
    section_executive_summary(doc)
    section_business_objective(doc)
    section_temporal_design(doc)
    section_data_architecture(doc)
    section_synthetic_data(doc)
    section_cohort_features(doc)
    section_models(doc)
    section_training(doc)
    section_evaluation(doc)
    section_hcp(doc)
    section_implementation(doc)
    section_outputs(doc)
    section_quickstart_evidence(doc, repo_root)
    section_limitations(doc)
    appendix_reference(doc)
    appendix_glossary(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="Repository root containing README.md and optional outputs/quickstart.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/Advanced_Therapy_Switch_Prediction_Project_Guide.docx"),
        help="DOCX destination relative to the current directory unless absolute.",
    )
    args = parser.parse_args()
    repo_root = args.repo_root.resolve()
    output_path = args.output if args.output.is_absolute() else (repo_root / args.output)
    build_document(repo_root, output_path.resolve())
    print(output_path.resolve())


if __name__ == "__main__":
    main()
